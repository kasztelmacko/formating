#!/usr/bin/env python3
"""
Example script demonstrating the first-occurrence-only numbering logic.
Shows how the system handles multiple consecutive paragraphs with the same style.
"""

from docx import Document
from document_formatter_config import DocumentFormatterConfig
from document_formatting_agent import DocumentFormattingAgent

def main():
    """
    Example usage of the first-occurrence-only numbering system.
    """
    # Load configuration
    config = DocumentFormatterConfig.load_and_validate_yaml(
        input_dir="data/input",
        style_filename="style_config.yaml",
        schema_filename="style_config_schema.yaml"
    )
    
    # Load your document
    doc_path = "data/input/test_yaml.docx"  # Update this path to your document
    doc = Document(doc_path)
    
    print("Original document paragraphs:")
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.style.name in ['chapter_titles', 'subchapter_titles_level_2', 'subchapter_titles_level_3']:
            print(f"  {i}: [{paragraph.style.name}] {paragraph.text}")
    
    # Apply formatting using the agent
    agent = DocumentFormattingAgent(doc, config)
    agent.apply_chapter_section_styles()
    
    print("\nAfter applying first-occurrence-only section numbering:")
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.style.name in ['chapter_titles', 'subchapter_titles_level_2', 'subchapter_titles_level_3']:
            print(f"  {i}: [{paragraph.style.name}] {paragraph.text}")
    
    # Save the modified document
    output_path = "data/output/test_with_first_occurrence_numbering.docx"
    doc.save(output_path)
    print(f"\nModified document saved to: {output_path}")

if __name__ == "__main__":
    main()
