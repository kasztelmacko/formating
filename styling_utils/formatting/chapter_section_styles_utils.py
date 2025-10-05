from docx.document import Document

from ..numbering.numbering_utils import (
    process_paragraph_text,
    update_paragraph_numbering,
)


def apply_chapter_page_breaks(
    doc: Document, style_names_mapping: dict[str, str]
) -> None:
    """
    Ensure only the first paragraph of each 'chapter_titles' block starts on a new page.
    """
    page_break_applied = False

    for paragraph in doc.paragraphs:
        if paragraph.style.name == style_names_mapping["chapter_titles"]:
            if not page_break_applied:
                paragraph.paragraph_format.page_break_before = True
                page_break_applied = True
        else:
            page_break_applied = False


def apply_chapter_section_numbering_format(
    doc: Document,
    style_definitions: dict[str, dict[str, str | dict[str, str]]],
    style_attributes_names_mapping: dict[str, str],
    chapter_section_numbering_regex: dict[str, str],
    renumbering_regex: dict[str, str] | None = None,
) -> None:
    """
    Adjust numbering in chapter/section titles based on YAML config.
    - numbering_format: { type: ROMAN|ARABIC, side: LEFT|RIGHT, separator: " " }
    """

    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name
        style_def = style_definitions.get(style_name)

        if not style_def:
            continue

        numbering_def = style_def.get(
            style_attributes_names_mapping["numbering_format"], {}
        )
        numbering_type = numbering_def.get("type")
        numbering_side = numbering_def.get("side")
        separator = numbering_def.get("separator", " ")

        if not numbering_type or not numbering_side:
            continue

        processed_text = process_paragraph_text(
            paragraph.text.strip(),
            numbering_type.upper(),
            numbering_side.upper(),
            chapter_section_numbering_regex,
            separator,
        )

        if processed_text != paragraph.text:
            paragraph.text = " ".join(processed_text.split())


def apply_section_numbering_order(
    doc: Document,
    style_names_mapping: dict[str, str],
    style_definitions: dict[str, dict[str, str | dict[str, str]]] | None = None,
    style_attributes_names_mapping: dict[str, str] | None = None,
    chapter_section_numbering_regex: dict[str, str] | None = None,
    renumbering_regex: dict[str, str] | None = None,
) -> None:
    """
    Adjust section numbering based on hierarchy:
    1. Find first paragraph with style chapter_titles and assign current_chapter = 1
    2. Find all subchapter_titles_level_2 until next chapter_titles and assign them
       current_chapter = current_chapter and subchapter_titles_level_2 grows by one
    3. Find all subchapter_titles_level_3 until next subchapter_titles_level_2 and assign them
       current_chapter = current_chapter, subchapter_titles_level_2 = subchapter_titles_level_2,
       and subchapter_titles_level_3 grows by one
    """

    current_chapter = 0
    current_subchapter_level_2 = 0
    current_subchapter_level_3 = 0

    in_chapter = False
    chapter_numbering_applied = False

    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name

        if style_name == style_names_mapping["chapter_titles"]:
            if not chapter_numbering_applied:
                current_chapter += 1
                current_subchapter_level_2 = 0
                current_subchapter_level_3 = 0
                in_chapter = True
                chapter_numbering_applied = True

                paragraph.text = update_paragraph_numbering(
                    paragraph.text,
                    current_chapter,
                    None,
                    None,
                    style_definitions,
                    style_attributes_names_mapping,
                    style_name,
                    chapter_section_numbering_regex,
                    renumbering_regex=renumbering_regex,
                )

        elif style_name == style_names_mapping["subchapter_titles_level_2"]:
            if in_chapter:
                current_subchapter_level_2 += 1
                current_subchapter_level_3 = 0

                paragraph.text = update_paragraph_numbering(
                    paragraph.text,
                    current_chapter,
                    current_subchapter_level_2,
                    None,
                    style_definitions,
                    style_attributes_names_mapping,
                    style_name,
                    chapter_section_numbering_regex,
                    renumbering_regex=renumbering_regex,
                )

        elif style_name == style_names_mapping["subchapter_titles_level_3"]:
            if in_chapter and current_subchapter_level_2 > 0:
                current_subchapter_level_3 += 1

                paragraph.text = update_paragraph_numbering(
                    paragraph.text,
                    current_chapter,
                    current_subchapter_level_2,
                    current_subchapter_level_3,
                    style_definitions,
                    style_attributes_names_mapping,
                    style_name,
                    chapter_section_numbering_regex,
                    renumbering_regex=renumbering_regex,
                )
        else:
            chapter_numbering_applied = False
