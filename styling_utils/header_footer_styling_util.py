from docx.oxml.shared import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Inches
import re
from style_mapping_config import HEADER_FOOTER_FIELD_MAPPINGS, HEADER_FOOTER_LAYOUT_CONFIG, FONT_MAPPING
from styling_utils.style_appliers import map_config_to_docx_attributes

def apply_header_footer_styles(doc, header_footer_config: dict, style_attributes_names_mapping: dict, 
                              field_mappings: dict, font_mapping: dict, layout_config: dict):
    """
    Apply header and footer styles and content to the document.
    
    Args:
        doc: The docx Document object
        header_footer_config: Configuration dictionary containing header/footer rules
        style_attributes_names_mapping: Mapping for style attribute names
        field_mappings: Field mapping configuration (defaults to HEADER_FOOTER_FIELD_MAPPINGS)
        font_mapping: Font mapping configuration (defaults to FONT_MAPPING)
        layout_config: Layout configuration (defaults to HEADER_FOOTER_LAYOUT_CONFIG)
    """
    section = doc.sections[0]

    header_style = header_footer_config.get("header_style", {})
    header_content = header_footer_config.get("header_content", {})
    
    if header_style or header_content:
        header = section.header
        _apply_header_footer_formatting(header, header_style, header_content, style_attributes_names_mapping,
                                       field_mappings, font_mapping, layout_config)

    footer_style = header_footer_config.get("footer_style", {})
    footer_content = header_footer_config.get("footer_content", {})
    
    if footer_style or footer_content:
        footer = section.footer
        _apply_header_footer_formatting(footer, footer_style, footer_content, style_attributes_names_mapping,
                                       field_mappings, font_mapping, layout_config)


def _apply_header_footer_formatting(header_footer, style_def: dict, content_def: dict, style_attributes_names_mapping: dict,
                                   field_mappings, font_mapping, layout_config):
    """
    Apply formatting and content to a header or footer using python-docx native methods.
    
    Args:
        header_footer: The header or footer object
        style_def: Style definition dictionary
        content_def: Content definition dictionary with left, center, right positions
        style_attributes_names_mapping: Mapping for style attribute names
        field_mappings: Field mapping configuration
        font_mapping: Font mapping configuration
        layout_config: Layout configuration
    """
    for paragraph in header_footer.paragraphs[:]:
        paragraph._element.getparent().remove(paragraph._element)
    
    content_positions = [
        content_def.get('left', ''),
        content_def.get('center', ''),
        content_def.get('right', '')
    ]
    
    non_empty_positions = [content for content in content_positions if content.strip()]
    
    if len(non_empty_positions) > 1 and layout_config['use_table_for_multiple_positions']:
        _create_table_layout_native(header_footer, style_def, content_def, style_attributes_names_mapping,
                                   field_mappings, font_mapping)
    else:
        _create_simple_layout_native(header_footer, style_def, content_def, style_attributes_names_mapping,
                                    field_mappings, font_mapping)


def _create_simple_layout_native(header_footer, style_def: dict, content_def: dict, style_attributes_names_mapping: dict,
                                field_mappings, font_mapping):
    """Create a simple paragraph-based layout using python-docx native methods."""
    content_mapping = [
        (content_def.get('center', ''), WD_ALIGN_PARAGRAPH.CENTER),
        (content_def.get('right', ''), WD_ALIGN_PARAGRAPH.RIGHT),
        (content_def.get('left', ''), WD_ALIGN_PARAGRAPH.LEFT)
    ]
    
    content_text, alignment = next(((text, align) for text, align in content_mapping if text.strip()), ('', None))
    
    if not content_text:
        return
    
    paragraph = _get_or_create_paragraph(header_footer)
    paragraph.alignment = alignment
    _add_formatted_content(paragraph, content_text, style_def, style_attributes_names_mapping,
                          field_mappings, font_mapping)


def _create_table_layout_native(header_footer, style_def: dict, content_def: dict, style_attributes_names_mapping: dict,
                               field_mappings, font_mapping):
    """Create a tab-based layout for multiple content positions."""
    paragraph = _get_or_create_paragraph(header_footer)
    _setup_tab_stops(paragraph)

    content_positions = [
        ('left', content_def.get('left', ''), False),
        ('center', content_def.get('center', ''), True), 
        ('right', content_def.get('right', ''), True)
    ]
    
    for position, content, needs_tab in content_positions:
        if content.strip():
            if needs_tab:
                paragraph.add_run('\t')
            _add_formatted_content(paragraph, content, style_def, style_attributes_names_mapping,
                                  field_mappings, font_mapping)
        elif needs_tab and _has_content_after(content_positions, position):
            paragraph.add_run('\t')


def _get_or_create_paragraph(header_footer):
    """Get the first paragraph or create a new one, clearing existing content."""
    if header_footer.paragraphs:
        paragraph = header_footer.paragraphs[0]
        paragraph.clear()
    else:
        paragraph = header_footer.add_paragraph()
    return paragraph


def _setup_tab_stops(paragraph):
    """Set up standard tab stops for header/footer layout."""
    tab_stops = paragraph.paragraph_format.tab_stops
    tab_stops.clear_all()
    tab_stops.add_tab_stop(Inches(3.25), WD_TAB_ALIGNMENT.CENTER)
    tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)


def _add_formatted_content(paragraph, content_text, style_def, style_attributes_names_mapping,
                          field_mappings, font_mapping):
    """Add formatted content to a paragraph."""
    run = paragraph.add_run()
    _apply_font_formatting_integrated(run, style_def, style_attributes_names_mapping, font_mapping)
    _add_content_with_fields_native(run, content_text, field_mappings)


def _has_content_after(content_positions, current_position):
    """Check if there's any content in positions after the current one."""
    current_index = next(i for i, (pos, _, _) in enumerate(content_positions) if pos == current_position)
    return any(content.strip() for _, content, _ in content_positions[current_index + 1:])


def _apply_font_formatting_integrated(run, style_def: dict, style_attributes_names_mapping: dict, font_mapping):
    """Apply font formatting using the existing style application system."""
    font_def = style_def.get(style_attributes_names_mapping.get("font_format", "font_format"), {})
    
    if not font_def:
        return

    map_config_to_docx_attributes(
        target=run.font,
        config_data=font_def,
        mapping=font_mapping
    )


def _add_content_with_fields_native(run, content_text, field_mappings):
    """Add content with field processing using a cleaner approach."""
    has_fields = any(re.search(pattern, content_text, re.IGNORECASE) for pattern, _ in field_mappings)
    if has_fields:
        _add_content_with_dynamic_fields(run._element, content_text, field_mappings)
    else:
        run.add_text(content_text)

def apply_header_footer_to_all_sections(doc, header_footer_config: dict, style_attributes_names_mapping: dict, field_mappings: dict, font_mapping: dict, layout_config: dict):
    """
    Apply header and footer styles to all sections in the document.
    
    Args:
        doc: The docx Document object
        header_footer_config: Configuration dictionary containing header/footer rules
        style_attributes_names_mapping: Mapping for style attribute names
    """
    for section in doc.sections:
        header_style = header_footer_config.get("header_style", {})
        header_content = header_footer_config.get("header_content", {})
        
        if header_style or header_content:
            header = section.header
            _apply_header_footer_formatting(header, header_style, header_content, style_attributes_names_mapping,
                                           field_mappings, font_mapping, layout_config)

        footer_style = header_footer_config.get("footer_style", {})
        footer_content = header_footer_config.get("footer_content", {})
        
        if footer_style or footer_content:
            footer = section.footer
            _apply_header_footer_formatting(footer, footer_style, footer_content, style_attributes_names_mapping,
                                           field_mappings, font_mapping, layout_config)


def _add_content_with_dynamic_fields(run_element, content_text, field_mappings):
    """
    Add content to a run element, processing dynamic field placeholders.
    
    Args:
        run_element: The XML run element to add content to
        content_text: Text content that may contain dynamic field placeholders
        field_mappings: Field mapping configuration
    """
    remaining_text = content_text
    
    while remaining_text:
        earliest_match = None
        earliest_pos = len(remaining_text)
        field_code = None
        
        for pattern, code in field_mappings:
            match = re.search(pattern, remaining_text, re.IGNORECASE)
            if match and match.start() < earliest_pos:
                earliest_match = match
                earliest_pos = match.start()
                field_code = code
        
        if earliest_match:
            before_text = remaining_text[:earliest_pos]
            if before_text:
                t_element = OxmlElement('w:t')
                if before_text.endswith(' ') or before_text.startswith(' '):
                    t_element.set(qn('xml:space'), 'preserve')
                t_element.text = before_text
                run_element.append(t_element)

            _add_field_code(run_element, field_code)

            remaining_text = remaining_text[earliest_match.end():]
        else:
            if remaining_text:
                t_element = OxmlElement('w:t')
                if remaining_text.endswith(' ') or remaining_text.startswith(' '):
                    t_element.set(qn('xml:space'), 'preserve')
                t_element.text = remaining_text
                run_element.append(t_element)
            break


def _add_field_code(run_element, field_code):
    """
    Add a Word field code to a run element using fldSimple approach.
    
    Args:
        run_element: The XML run element to add the field to
        field_code: The Word field code (e.g., 'PAGE', 'NUMPAGES', etc.)
    """
    fld_simple = OxmlElement('w:fldSimple')
    fld_simple.set(qn('w:instr'), field_code)
    run_element.append(fld_simple)