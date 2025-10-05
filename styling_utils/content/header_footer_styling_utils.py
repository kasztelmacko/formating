import re
from typing import Callable

from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.shared import qn
from docx.section import FooterPart, HeaderPart
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from styling_utils.core.style_appliers import map_config_to_docx_attributes


def apply_header_footer_styles(
    doc: Document,
    header_footer_config: dict[str, dict[str, str | dict[str, str]]],
    style_attributes_names_mapping: dict[str, str],
    field_mappings: list[tuple[str, str]],
    font_mapping: dict[str, tuple[str, Callable | None]],
    layout_config: dict[str, bool],
) -> None:
    """Apply header and footer styles and content to the document."""
    section = doc.sections[0]

    header_style = header_footer_config.get("header_style", {})
    header_content = header_footer_config.get("header_content", {})

    if header_style or header_content:
        header = section.header
        _apply_header_footer_formatting(
            header,
            header_style,
            header_content,
            style_attributes_names_mapping,
            field_mappings,
            font_mapping,
            layout_config,
        )

    footer_style = header_footer_config.get("footer_style", {})
    footer_content = header_footer_config.get("footer_content", {})

    if footer_style or footer_content:
        footer = section.footer
        _apply_header_footer_formatting(
            footer,
            footer_style,
            footer_content,
            style_attributes_names_mapping,
            field_mappings,
            font_mapping,
            layout_config,
        )


def _apply_header_footer_formatting(
    header_footer: HeaderPart | FooterPart,
    style_def: dict[str, str | dict[str, str]],
    content_def: dict[str, str],
    style_attributes_names_mapping: dict[str, str],
    field_mappings: list[tuple[str, str]],
    font_mapping: dict[str, tuple[str, Callable | None]],
    layout_config: dict[str, bool],
) -> None:
    """Apply formatting and content to a header or footer using python-docx native methods."""
    for paragraph in header_footer.paragraphs[:]:
        paragraph._element.getparent().remove(paragraph._element)

    content_positions = [
        content_def.get("left", ""),
        content_def.get("center", ""),
        content_def.get("right", ""),
    ]

    non_empty_positions = [content for content in content_positions if content.strip()]

    if (
        len(non_empty_positions) > 1
        and layout_config["use_table_for_multiple_positions"]
    ):
        _create_table_layout_native(
            header_footer,
            style_def,
            content_def,
            style_attributes_names_mapping,
            field_mappings,
            font_mapping,
        )
    else:
        _create_simple_layout_native(
            header_footer,
            style_def,
            content_def,
            style_attributes_names_mapping,
            field_mappings,
            font_mapping,
        )


def _create_simple_layout_native(
    header_footer: HeaderPart | FooterPart,
    style_def: dict[str, str | dict[str, str]],
    content_def: dict[str, str],
    style_attributes_names_mapping: dict[str, str],
    field_mappings: list[tuple[str, str]],
    font_mapping: dict[str, tuple[str, Callable | None]],
) -> None:
    """Create a simple paragraph-based layout using python-docx native methods."""
    content_mapping = [
        (content_def.get("center", ""), WD_ALIGN_PARAGRAPH.CENTER),
        (content_def.get("right", ""), WD_ALIGN_PARAGRAPH.RIGHT),
        (content_def.get("left", ""), WD_ALIGN_PARAGRAPH.LEFT),
    ]

    content_text, alignment = next(
        ((text, align) for text, align in content_mapping if text.strip()), ("", None)
    )

    if not content_text:
        return

    paragraph = _get_or_create_paragraph(header_footer)
    paragraph.alignment = alignment
    _add_formatted_content(
        paragraph,
        content_text,
        style_def,
        style_attributes_names_mapping,
        field_mappings,
        font_mapping,
    )


def _create_table_layout_native(
    header_footer: HeaderPart | FooterPart,
    style_def: dict[str, str | dict[str, str]],
    content_def: dict[str, str],
    style_attributes_names_mapping: dict[str, str],
    field_mappings: list[tuple[str, str]],
    font_mapping: dict[str, tuple[str, Callable | None]],
) -> None:
    """Create a tab-based layout for multiple content positions."""
    paragraph = _get_or_create_paragraph(header_footer)
    _setup_tab_stops(paragraph)

    content_positions = [
        ("left", content_def.get("left", ""), False),
        ("center", content_def.get("center", ""), True),
        ("right", content_def.get("right", ""), True),
    ]

    for position, content, needs_tab in content_positions:
        if content.strip():
            if needs_tab:
                paragraph.add_run("\t")
            _add_formatted_content(
                paragraph,
                content,
                style_def,
                style_attributes_names_mapping,
                field_mappings,
                font_mapping,
            )
        elif needs_tab and _has_content_after(content_positions, position):
            paragraph.add_run("\t")


def _get_or_create_paragraph(
    header_footer: HeaderPart | FooterPart,
) -> Paragraph:
    """Get the first paragraph or create a new one, clearing existing content."""
    if header_footer.paragraphs:
        paragraph = header_footer.paragraphs[0]
        paragraph.clear()
    else:
        paragraph = header_footer.add_paragraph()
    return paragraph


def _setup_tab_stops(paragraph: Paragraph) -> None:
    """Set up standard tab stops for header/footer layout."""
    tab_stops = paragraph.paragraph_format.tab_stops
    tab_stops.clear_all()
    tab_stops.add_tab_stop(Inches(3.25), WD_TAB_ALIGNMENT.CENTER)
    tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)


def _add_formatted_content(
    paragraph: Paragraph,
    content_text: str,
    style_def: dict[str, str | dict[str, str]],
    style_attributes_names_mapping: dict[str, str],
    field_mappings: list[tuple[str, str]],
    font_mapping: dict[str, tuple[str, Callable | None]],
) -> None:
    """Add formatted content to a paragraph."""
    run = paragraph.add_run()
    _apply_font_formatting_integrated(
        run, style_def, style_attributes_names_mapping, font_mapping
    )
    _add_content_with_fields_native(run, content_text, field_mappings)


def _has_content_after(
    content_positions: list[tuple[str, str, bool]], current_position: str
) -> bool:
    """Check if there's any content in positions after the current one."""
    current_index = next(
        i for i, (pos, _, _) in enumerate(content_positions) if pos == current_position
    )
    return any(
        content.strip() for _, content, _ in content_positions[current_index + 1 :]
    )


def _apply_font_formatting_integrated(
    run: Run,
    style_def: dict[str, str | dict[str, str]],
    style_attributes_names_mapping: dict[str, str],
    font_mapping: dict[str, tuple[str, Callable | None]],
) -> None:
    """Apply font formatting using the existing style application system."""
    font_def = style_def.get(
        style_attributes_names_mapping.get("font_format", "font_format"), {}
    )

    if not font_def:
        return

    map_config_to_docx_attributes(
        target=run.font, config_data=font_def, mapping=font_mapping
    )


def _add_content_with_fields_native(
    run: Run, content_text: str, field_mappings: list[tuple[str, str]]
) -> None:
    """Add content with field processing using a cleaner approach."""
    has_fields = any(
        re.search(pattern, content_text, re.IGNORECASE) for pattern, _ in field_mappings
    )
    if has_fields:
        _add_content_with_dynamic_fields(run._element, content_text, field_mappings)
    else:
        run.add_text(content_text)


def apply_header_footer_to_all_sections(
    doc: Document,
    header_footer_config: dict[str, dict[str, str | dict[str, str]]],
    style_attributes_names_mapping: dict[str, str],
    field_mappings: list[tuple[str, str]],
    font_mapping: dict[str, tuple[str, Callable | None]],
    layout_config: dict[str, bool],
) -> None:
    """Apply header and footer styles to all sections in the document."""
    for section in doc.sections:
        header_style = header_footer_config.get("header_style", {})
        header_content = header_footer_config.get("header_content", {})

        if header_style or header_content:
            header = section.header
            _apply_header_footer_formatting(
                header,
                header_style,
                header_content,
                style_attributes_names_mapping,
                field_mappings,
                font_mapping,
                layout_config,
            )

        footer_style = header_footer_config.get("footer_style", {})
        footer_content = header_footer_config.get("footer_content", {})

        if footer_style or footer_content:
            footer = section.footer
            _apply_header_footer_formatting(
                footer,
                footer_style,
                footer_content,
                style_attributes_names_mapping,
                field_mappings,
                font_mapping,
                layout_config,
            )


def _add_content_with_dynamic_fields(
    run_element: OxmlElement, content_text: str, field_mappings: list[tuple[str, str]]
) -> None:
    """Add content to a run element, processing dynamic field placeholders."""
    remaining_text = content_text

    while remaining_text:
        earliest_match = None
        earliest_pos = len(remaining_text)
        field_code = None

        for pattern, code in field_mappings:
            match = re.search(pattern, remaining_text, re.IGNORECASE)
            if match and match.start() < earliest_pos:
                earliest_match = match
                earliest_pos = match.start()
                field_code = code

        if earliest_match:
            before_text = remaining_text[:earliest_pos]
            if before_text:
                t_element = OxmlElement("w:t")
                if before_text.endswith(" ") or before_text.startswith(" "):
                    t_element.set(qn("xml:space"), "preserve")
                t_element.text = before_text
                run_element.append(t_element)

            _add_field_code(run_element, field_code)

            remaining_text = remaining_text[earliest_match.end() :]
        else:
            if remaining_text:
                t_element = OxmlElement("w:t")
                if remaining_text.endswith(" ") or remaining_text.startswith(" "):
                    t_element.set(qn("xml:space"), "preserve")
                t_element.text = remaining_text
                run_element.append(t_element)
            break


def _add_field_code(run_element: OxmlElement, field_code: str) -> None:
    """Add a Word field code to a run element using fldSimple approach."""
    fld_simple = OxmlElement("w:fldSimple")
    fld_simple.set(qn("w:instr"), field_code)
    run_element.append(fld_simple)
